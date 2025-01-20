import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

def combine_core_files_to_word(directory_path, output_doc_name='project_documentation.docx'):
    # Create Word document
    doc = Document()
    
    # Define target files
    core_files = {
        'URLs': '*urls.py',
        'Views': '*views.py',
        'Models': '*models.py'
    }
    
    # Style settings
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        heading.style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
    def add_code(text):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    
    # Process each file type
    for category, pattern in core_files.items():
        add_heading(f"{category} Documentation", 1)
        
        # Find all matching files
        files = list(Path(directory_path).rglob(pattern))
        files.sort()  # Sort files for consistent output
        
        for file_path in files:
            try:
                # Add file path as subheading
                add_heading(f"File: {file_path.relative_to(directory_path)}", 2)
                
                # Read and add content
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    add_code(content)
                
                # Add spacing between files
                doc.add_paragraph()
                
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
    
    # Save the document
    try:
        doc.save(output_doc_name)
        print(f"Word document successfully created: {output_doc_name}")
    except Exception as e:
        print(f"Error saving document: {str(e)}")

if __name__ == "__main__":
    # Get current directory
    current_dir = os.getcwd()
    
    # Call the function
    combine_core_files_to_word(current_dir)